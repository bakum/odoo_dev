import re
from langdetect import detect

from odoo import models, fields, api
import io, base64
import fitz  # PyMuPDF для PDF
import docx  # Для DOCX

from ..services.chunking_service import ChunkingService
from ..services.index_service import RagIndexService  # импорт нового сервиса
from ..services.triple_extractor_service import TripleExtractorService


class RagDocument(models.Model):
    _name = "rag.document"
    _description = "RAG Document"

    name = fields.Char(required=True)
    file = fields.Binary("File", required=True, attachment=True)
    filename = fields.Char()
    file_type = fields.Selection([
        ('pdf', 'PDF'),
        ('docx', 'DOCX'),
        ('other', 'Other')
    ], string="File Type")
    chunk_ids = fields.One2many("rag.chunk", "document_id", string="Chunks")
    language = fields.Char("Language")
    preview = fields.Text("Preview", compute="_compute_preview")

    @api.model
    def _extract_knowledge_graph(self):
        self.ensure_one()
        service = TripleExtractorService()
        triples = service.extract_triples(self.chunk_ids)
        for subj_name, relation, obj_name in triples:
            subj = self.env['rag.entity']._get_or_create(subj_name.strip())
            obj = self.env['rag.entity']._get_or_create(obj_name.strip())

            self.env['rag.relation'].create({
                'subject_id': subj.id,
                'object_id': obj.id,
                'relation': relation.strip(),
                'document_id': self.id,
            })

    @api.onchange('file', 'name')
    def _compute_filename(self):
        for rec in self:
            if rec.file and rec.name:
                if not rec.filename:
                    rec.filename = f"{rec.name}"
            else:
                rec.filename = ''

    @api.model
    def create(self, vals):
        if 'filename' in vals:
            ext = vals['filename'].split('.')[-1].lower()
            vals['file_type'] = ext if ext in ['pdf', 'docx'] else 'other'
        return super().create(vals)

    def action_reindex(self):
        self._extract_chunks()
        index_service = RagIndexService(self.env)
        index_service.compute_embeddings()
        index_service.build_index()

    def _extract_chunks(self, chunk_size=5):
        for rec in self:
            rec.chunk_ids.unlink()
            if not rec.file or not rec.filename:
                continue

            extension = rec.filename.split('.')[-1].lower()
            data = io.BytesIO(base64.b64decode(rec.file))

            if extension == 'pdf':
                doc = fitz.open(stream=data.read(), filetype="pdf")
                for page_num, page in enumerate(doc, start=1):
                    text = page.get_text()
                    self._create_chunks_from_text(rec, text, page_num, chunk_size)

            elif extension == 'docx':
                document = docx.Document(data)
                full_text = "\n".join([para.text for para in document.paragraphs])
                self._create_chunks_from_text(rec, full_text, page_num=1, chunk_size=chunk_size)

            else:
                raise ValueError(f"Unsupported file type: {extension}")

            try:
                rec.language = detect(full_text if extension == 'docx' else text)
            except Exception:
                rec.language = 'unknown'

    def _create_chunks_from_text(self, rec, text, page_num, chunk_size=5):
        chunker = ChunkingService(max_tokens=300, overlap=50)
        chunks = chunker.split_text(text)

        for idx, chunk in enumerate(chunks):
            self.env['rag.chunk'].create({
                'document_id': rec.id,
                'content': chunk['text'],
                'position': idx,
                'page_number': page_num,
                'char_start': chunk['start'],
                'char_end': chunk['end'],
            })

    @api.depends('file', 'filename')
    def _compute_preview(self):
        for rec in self:
            if not rec.file or not rec.filename or not rec.filename.endswith('.docx'):
                rec.preview = ''
                continue
            data = io.BytesIO(base64.b64decode(rec.file))
            try:
                document = docx.Document(data)
                rec.preview = "\n".join([p.text for p in document.paragraphs if p.text.strip()])
            except Exception:
                rec.preview = "Ошибка чтения документа."

    def action_reindex_all(self):
        """Переиндексация всех документов: пересоздание эмбеддингов и FAISS-индекса"""
        index_service = RagIndexService(self.env)
        index_service.compute_embeddings()
        index_service.build_index()

    def action_extract_knowledge(self):
        for doc in self:
            doc._extract_knowledge_graph()
