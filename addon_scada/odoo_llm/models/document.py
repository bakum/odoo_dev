import io
import base64
import PyPDF2
from docx import Document as DocxDocument
import logging

from odoo import api, fields, models

_logger = logging.getLogger(__name__)

class Document(models.Model):
    _name = "llm.document"
    _description = "Corporative Document"

    name      = fields.Char("Name", required=True)
    file      = fields.Binary("File", required=True)
    text_content      = fields.Text("Text")
    embedding = fields.Text("Embedding")

    def action_extract_text(self):
        for rec in self:
            if rec.file:
                rec.text_content = self._extract_text_from_binary(rec.file, rec.name)
                vec = self.env['llm.embedding_service'].embed(rec.text_content)
                rec.embedding = self.env['llm.embedding_service'].serialize(vec)

    def build_index(self):
        """Build FAISS index for all documents."""
        return self.env['llm.vector.service'].index_all()

    def _extract_text_from_binary(self, data, fname):

        if not isinstance(fname, str) or '.' not in fname:
            return ""

        raw = base64.b64decode(data or b"")
        ext = fname.split('.')[-1].lower()
        if ext == 'pdf':
            reader = PyPDF2.PdfReader(io.BytesIO(raw))
            return "\n".join(page.extract_text() or '' for page in reader.pages)
        if ext in ('docx', 'doc'):
            doc = DocxDocument(io.BytesIO(raw))
            return "\n".join(p.text for p in doc.paragraphs if p.text)
        return ""

    # 🚨 Автоматическая переиндексация после создания
    @api.model
    def create(self, vals):
        rec = super().create(vals)

        # Генерация эмбеддинга, если нужно
        if rec.text_content:
            emb_service = self.env['llm.embedding_service']
            vec = emb_service.embed(rec.text_content)
            rec.embedding = emb_service.serialize(vec)

        # Переиндексация
        self.env['llm.vector.index'].build_index()

        return rec

    # 🚨 Автоматическая переиндексация после удаления
    def unlink(self):
        res = super().unlink()
        self.env['llm.vector.index'].build_index()
        return res